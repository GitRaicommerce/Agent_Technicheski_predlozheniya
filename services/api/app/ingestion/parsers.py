"""
Парсъри за PDF, DOCX и Markdown.
Използва markitdown (Microsoft) като primary метод — запазва структурата
(заглавия, таблици, списъци) в Markdown формат преди chunking.
Fallback към директно pypdf/python-docx при грешка.
"""

from __future__ import annotations

import io
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any


_AUDIT_SCHEMA_VERSION = 1
_MIN_PAGE_TEXT_CHARS = 80
_MIN_MARKDOWN_REFERENCE_COVERAGE = 0.75
_OPEN_DATALOADER_PAGE_SEPARATOR = "\n\n<!-- Page %page-number% -->\n\n"


def extract_chunks(content: bytes, filename: str) -> list[dict[str, Any]]:
    """
    Извлича текстови chunks от PDF, DOCX или Markdown файл.
    Опитва markitdown first (по-добра структурна вярност),
    след това fallback към нативните парсъри.
    Връща list от dict с полета: type, text, page, section_path.
    """
    chunks, _report = extract_chunks_with_audit(content, filename)
    return chunks

    filename_lower = filename.lower()

    # Markdown — директно chunking без конвертиране
    if filename_lower.endswith((".md", ".txt")):
        text = content.decode("utf-8", errors="replace")
        return _chunks_from_markdown(text)

    # PDF/DOCX — опитай markitdown първо
    md_text = _to_markdown_via_markitdown(content, filename)
    if md_text:
        return _chunks_from_markdown(md_text)

    # Fallback към нативните парсъри
    if filename_lower.endswith(".pdf"):
        return _extract_pdf(content)
    elif filename_lower.endswith(".docx"):
        return _extract_docx(content)
    else:
        return _extract_ocr(content)


def extract_chunks_with_audit(
    content: bytes, filename: str
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """
    Extract chunks and return a compact QA report for ingestion controls.

    The report stores metrics and flags only, not full extracted text.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith((".md", ".txt")):
        text = content.decode("utf-8", errors="replace")
        chunks = _chunks_from_markdown(text)
        return chunks, _build_report(
            filename=filename,
            file_type="text",
            chunks=chunks,
            primary_method="plain_text",
            methods_attempted=["plain_text"],
        )

    if filename_lower.endswith(".pdf"):
        return _extract_pdf_with_audit(content, filename)

    methods_attempted = ["markitdown"]
    warnings: list[str] = []
    md_text = _to_markdown_via_markitdown(content, filename)
    if md_text:
        chunks = _chunks_from_markdown(md_text)
        return chunks, _build_report(
            filename=filename,
            file_type=_file_type(filename_lower),
            chunks=chunks,
            primary_method="markitdown",
            methods_attempted=methods_attempted,
        )

    warnings.append("markitdown_returned_no_text")
    if filename_lower.endswith(".docx"):
        methods_attempted.append("python_docx")
        chunks = _extract_docx(content)
        primary_method = "python_docx"
    else:
        methods_attempted.append("ocr")
        chunks = _extract_ocr(content)
        primary_method = "ocr"

    return chunks, _build_report(
        filename=filename,
        file_type=_file_type(filename_lower),
        chunks=chunks,
        primary_method=primary_method,
        methods_attempted=methods_attempted,
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# markitdown conversion
# ---------------------------------------------------------------------------

def _to_markdown_via_markitdown(content: bytes, filename: str) -> str:
    """
    Конвертира файл до Markdown с Microsoft markitdown.
    Запазва заглавия, таблици и списъци.
    Връща празен стринг при грешка (→ fallback).
    """
    try:
        from markitdown import MarkItDown
        import tempfile, os

        # markitdown работи с файлове на диска
        suffix = "." + filename.rsplit(".", 1)[-1] if "." in filename else ".bin"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            md = MarkItDown()
            result = md.convert(tmp_path)
            return result.text_content or ""
        finally:
            os.unlink(tmp_path)
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Chunking from Markdown
# ---------------------------------------------------------------------------

def _chunks_from_markdown(md_text: str) -> list[dict[str, Any]]:
    """
    Разделя Markdown текст на смислени chunks, запазвайки section_path.
    Заглавия (# / ##) стават тип 'heading' и определят section_path.
    Таблици, списъци и абзаци → тип 'text'.
    """
    chunks: list[dict[str, Any]] = []
    current_section: str | None = None
    buffer_lines: list[str] = []
    page_counter = 0  # estimated — Markdown няма реални страници

    def _flush(section: str | None) -> None:
        nonlocal page_counter
        text = "\n".join(buffer_lines).strip()
        if len(text) >= 20:
            chunks.append({
                "type": "text",
                "text": text,
                "page": page_counter or None,
                "section_path": section,
            })
        buffer_lines.clear()

    lines = md_text.splitlines()
    i = 0
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Page break hint (markitdown вмъква \f или <!-- Page --> при PDF)
        if stripped in ("\f", "<!-- Page -->") or stripped.startswith("<!-- Page "):
            _flush(current_section)
            page_counter += 1
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            _flush(current_section)
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            current_section = title
            chunks.append({
                "type": "heading",
                "text": title,
                "page": page_counter or None,
                "section_path": title,
            })
            i += 1
            continue

        # Tables — collect entire table as one chunk
        if stripped.startswith("|"):
            if not in_table:
                _flush(current_section)
                in_table = True
            buffer_lines.append(line)
            i += 1
            # Check if next line continues the table
            if i < len(lines) and lines[i].strip().startswith("|"):
                continue
            else:
                in_table = False
                _flush(current_section)
            continue

        # Blank line → flush current paragraph
        if not stripped:
            _flush(current_section)
            i += 1
            continue

        # Regular text / list items
        buffer_lines.append(stripped)
        i += 1

    _flush(current_section)
    return chunks


def _extract_pdf_with_audit(
    content: bytes, filename: str
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    methods_attempted = ["opendataloader_pdf", "markitdown"]
    warnings: list[str] = []

    opendataloader_text, opendataloader_warnings = _extract_pdf_via_opendataloader_markdown(
        content, filename
    )
    warnings.extend(opendataloader_warnings)
    opendataloader_chunks = (
        _annotate_chunks(
            _chunks_from_markdown(opendataloader_text),
            parser_method="opendataloader_pdf",
        )
        if opendataloader_text
        else []
    )
    if not opendataloader_text and not opendataloader_warnings:
        warnings.append("opendataloader_returned_no_text")

    md_text = _to_markdown_via_markitdown(content, filename)
    md_chunks = (
        _annotate_chunks(
            _chunks_from_markdown(md_text),
            parser_method="markitdown",
        )
        if md_text
        else []
    )
    if not md_text:
        warnings.append("markitdown_returned_no_text")

    page_audits, page_texts, page_methods, page_warnings = _audit_pdf_pages(content)
    methods_attempted.extend(page_methods)
    warnings.extend(page_warnings)

    reference_chars = sum(len(text.strip()) for _page, text in page_texts)
    opendataloader_chars = _chunks_text_chars(opendataloader_chunks)
    markitdown_chars = _chunks_text_chars(md_chunks)

    chunks: list[dict[str, Any]] = []
    primary_method = "pdf_page_text"
    markdown_chars: int | None = None

    if _has_sufficient_reference_coverage(opendataloader_chars, reference_chars):
        chunks = opendataloader_chunks
        primary_method = "opendataloader_pdf"
        markdown_chars = opendataloader_chars
    elif _has_sufficient_reference_coverage(markitdown_chars, reference_chars):
        if opendataloader_chunks:
            warnings.append("opendataloader_low_reference_coverage")
        chunks = md_chunks
        primary_method = "markitdown"
        markdown_chars = markitdown_chars
    elif (
        reference_chars
        and opendataloader_chars < int(reference_chars * _MIN_MARKDOWN_REFERENCE_COVERAGE)
        and markitdown_chars < int(reference_chars * _MIN_MARKDOWN_REFERENCE_COVERAGE)
    ):
        warnings.append("used_page_text_fallback_low_markdown_coverage")
        chunks = _chunks_from_page_texts(page_texts)
    elif opendataloader_chars >= markitdown_chars and opendataloader_chunks:
        warnings.append("used_opendataloader_best_effort")
        chunks = opendataloader_chunks
        primary_method = "opendataloader_pdf"
        markdown_chars = opendataloader_chars
    elif md_chunks:
        warnings.append("used_markitdown_best_effort")
        chunks = md_chunks
        primary_method = "markitdown"
        markdown_chars = markitdown_chars
    else:
        warnings.append("used_page_text_fallback_no_markdown_chunks")
        chunks = _chunks_from_page_texts(page_texts)

    if primary_method != "pdf_page_text" and page_audits and not any(
        chunk.get("page") for chunk in chunks
    ):
        warnings.append("markdown_chunks_have_no_page_numbers")

    return chunks, _build_report(
        filename=filename,
        file_type="pdf",
        chunks=chunks,
        primary_method=primary_method,
        methods_attempted=methods_attempted,
        pages=page_audits,
        warnings=warnings,
        reference_chars=reference_chars,
        markdown_chars=markdown_chars,
    )


def _extract_pdf_via_opendataloader_markdown(
    content: bytes, filename: str
) -> tuple[str, list[str]]:
    try:
        import opendataloader_pdf
    except Exception:
        return "", ["opendataloader_not_installed"]

    suffix = Path(filename).suffix or ".pdf"
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / f"source{suffix}"
            output_dir = Path(tmpdir) / "output"
            input_path.write_bytes(content)
            output_dir.mkdir(parents=True, exist_ok=True)

            opendataloader_pdf.convert(
                input_path=[str(input_path)],
                output_dir=str(output_dir),
                format="markdown",
                quiet=True,
                keep_line_breaks=True,
                use_struct_tree=True,
                markdown_page_separator=_OPEN_DATALOADER_PAGE_SEPARATOR,
            )

            markdown_files = [
                *output_dir.rglob("*.md"),
                *output_dir.rglob("*.markdown"),
            ]
            if not markdown_files:
                return "", ["opendataloader_no_markdown_output"]

            markdown_path = sorted(markdown_files, key=lambda path: len(path.parts))[0]
            return markdown_path.read_text(encoding="utf-8", errors="replace"), []
    except Exception as exc:
        return "", [_classify_opendataloader_error(exc)]


def _annotate_chunks(
    chunks: list[dict[str, Any]], *, parser_method: str
) -> list[dict[str, Any]]:
    annotated: list[dict[str, Any]] = []
    for chunk in chunks:
        meta = dict(chunk.get("meta") or {})
        meta.setdefault("parser_method", parser_method)
        annotated.append(
            {
                **chunk,
                "parser_method": parser_method,
                "meta": meta,
            }
        )
    return annotated


def _has_sufficient_reference_coverage(chars: int, reference_chars: int) -> bool:
    if chars <= 0:
        return False
    if reference_chars <= 0:
        return True
    return chars >= int(reference_chars * _MIN_MARKDOWN_REFERENCE_COVERAGE)


def _classify_opendataloader_error(exc: Exception) -> str:
    if isinstance(exc, subprocess.CalledProcessError):
        stderr = str(exc.stderr or "")
        if "UnsupportedClassVersionError" in stderr:
            return "opendataloader_java_too_old"
        return f"opendataloader_cli_failed:{exc.returncode}"

    return f"opendataloader_failed:{type(exc).__name__}"


def _audit_pdf_pages(
    content: bytes,
) -> tuple[list[dict[str, Any]], list[tuple[int, str]], list[str], list[str]]:
    from pypdf import PdfReader

    methods_attempted = ["pypdf"]
    warnings: list[str] = []
    page_audits: list[dict[str, Any]] = []
    page_texts: list[tuple[int, str]] = []

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        return [], [], methods_attempted, [f"pdf_reader_failed:{type(exc).__name__}"]

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        method = "pypdf"
        flags: list[str] = []

        if _needs_ocr(text):
            ocr_text, ocr_method = _ocr_pdf_page(content, page_num, page)
            methods_attempted.append(ocr_method)
            if ocr_text.strip() and len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
                method = ocr_method
                flags.append("ocr_used")
            else:
                flags.append("ocr_unavailable_or_not_better")

        audit = _page_audit(page_num, method, text, flags)
        page_audits.append(audit)
        page_texts.append((page_num, text))

    return page_audits, page_texts, methods_attempted, warnings


def _ocr_pdf_page(content: bytes, page_num: int, page: Any) -> tuple[str, str]:
    rendered = _ocr_pdf_page_via_pdftoppm(content, page_num)
    if rendered.strip():
        return rendered, "ocr_pdftoppm"

    embedded = _ocr_page_image(page)
    if embedded.strip():
        return embedded, "ocr_embedded_image"

    return "", "ocr_unavailable"


def _ocr_pdf_page_via_pdftoppm(content: bytes, page_num: int) -> str:
    if not shutil.which("pdftoppm"):
        return ""

    try:
        import pytesseract
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = os.path.join(tmpdir, "source.pdf")
            out_prefix = os.path.join(tmpdir, "page")
            with open(pdf_path, "wb") as handle:
                handle.write(content)

            subprocess.run(
                [
                    "pdftoppm",
                    "-f",
                    str(page_num),
                    "-l",
                    str(page_num),
                    "-r",
                    "300",
                    "-png",
                    pdf_path,
                    out_prefix,
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=45,
            )

            images = [
                os.path.join(tmpdir, name)
                for name in os.listdir(tmpdir)
                if name.startswith("page") and name.endswith(".png")
            ]
            if not images:
                return ""
            with Image.open(images[0]) as image:
                return pytesseract.image_to_string(image, lang="bul+eng")
    except Exception:
        return ""


def _needs_ocr(text: str) -> bool:
    return len(text.strip()) < _MIN_PAGE_TEXT_CHARS


def _page_audit(
    page_num: int, method: str, text: str, flags: list[str] | None = None
) -> dict[str, Any]:
    clean = text.strip()
    words = re.findall(r"\w+", clean, flags=re.UNICODE)
    non_space = [ch for ch in clean if not ch.isspace()]
    readable = sum(1 for ch in non_space if ch.isalnum() or ch in ".,;:!?()[]{}-/+%")
    cyrillic = sum(1 for ch in clean if "\u0400" <= ch <= "\u04ff")
    latin = sum(1 for ch in clean if ("a" <= ch.lower() <= "z"))
    issues = list(flags or [])

    if not clean:
        issues.append("empty_text")
    elif len(clean) < _MIN_PAGE_TEXT_CHARS:
        issues.append("low_text_density")

    if non_space and readable / len(non_space) < 0.55:
        issues.append("possible_garbled_text")
    if clean and len(words) < 5:
        issues.append("very_few_words")

    return {
        "page": page_num,
        "method": method,
        "text_chars": len(clean),
        "word_count": len(words),
        "cyrillic_chars": cyrillic,
        "latin_chars": latin,
        "issues": sorted(set(issues)),
    }


def _chunks_from_page_texts(page_texts: list[tuple[int, str]]) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    for page_num, text in page_texts:
        for para in _split_paragraphs(text):
            if para.strip():
                chunks.append(
                    {
                        "type": "text",
                        "text": para.strip(),
                        "page": page_num,
                        "section_path": None,
                    }
                )
    return chunks


def _build_report(
    *,
    filename: str,
    file_type: str,
    chunks: list[dict[str, Any]],
    primary_method: str,
    methods_attempted: list[str],
    pages: list[dict[str, Any]] | None = None,
    warnings: list[str] | None = None,
    errors: list[str] | None = None,
    reference_chars: int | None = None,
    markdown_chars: int | None = None,
) -> dict[str, Any]:
    pages = pages or []
    warnings = list(warnings or [])
    errors = list(errors or [])

    page_issue_count = sum(1 for page in pages if page.get("issues"))
    if not chunks:
        errors.append("no_chunks_extracted")

    quality_status = "ok"
    if errors:
        quality_status = "error"
    elif warnings or page_issue_count:
        quality_status = "warning"

    return {
        "schema_version": _AUDIT_SCHEMA_VERSION,
        "filename": filename,
        "file_type": file_type,
        "quality_status": quality_status,
        "primary_method": primary_method,
        "methods_attempted": sorted(set(methods_attempted)),
        "page_count": len(pages) or None,
        "pages_with_text": sum(1 for page in pages if page.get("text_chars", 0) > 0)
        or None,
        "page_issue_count": page_issue_count,
        "chunk_count": len(chunks),
        "extracted_chars": _chunks_text_chars(chunks),
        "reference_chars": reference_chars,
        "markdown_chars": markdown_chars,
        "warnings": sorted(set(warnings)),
        "errors": sorted(set(errors)),
        "pages": pages,
    }


def _chunks_text_chars(chunks: list[dict[str, Any]]) -> int:
    return sum(len(str(chunk.get("text", "")).strip()) for chunk in chunks)


def _file_type(filename_lower: str) -> str:
    if filename_lower.endswith(".docx"):
        return "docx"
    if filename_lower.endswith(".doc"):
        return "doc"
    if filename_lower.endswith((".md", ".txt")):
        return "text"
    return "unknown"


# ---------------------------------------------------------------------------
# Fallback parsers (kept for reliability)
# ---------------------------------------------------------------------------

def _extract_pdf(content: bytes) -> list[dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    chunks = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            text = _ocr_page_image(page)

        for para in _split_paragraphs(text):
            if para.strip():
                chunks.append({
                    "type": "text",
                    "text": para.strip(),
                    "page": page_num,
                    "section_path": None,
                })
    return chunks


def _extract_docx(content: bytes) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(io.BytesIO(content))
    chunks = []
    current_section: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower()
        chunk_type = "heading" if "heading" in style_name else "text"
        section_path = (
            text if chunk_type == "heading"
            else (current_section[-1] if current_section else None)
        )

        if chunk_type == "heading":
            current_section = [text]

        chunks.append({
            "type": chunk_type,
            "text": text,
            "page": None,
            "section_path": section_path,
        })

    for table in doc.tables:
        section_path = current_section[-1] if current_section else None
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                chunks.append({
                    "type": "table_row",
                    "text": " | ".join(cell_texts),
                    "page": None,
                    "section_path": section_path,
                })
    return chunks


def _extract_ocr(content: bytes) -> list[dict[str, Any]]:
    """OCR fallback с Tesseract (Bulgarian + English)."""
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image, lang="bul+eng")
        return [
            {"type": "text", "text": para.strip(), "page": 1, "section_path": None}
            for para in _split_paragraphs(text)
            if para.strip()
        ]
    except Exception:
        return []


def _ocr_page_image(page) -> str:
    """OCR на PDF страница при липса на текстов слой."""
    try:
        import pytesseract
        from PIL import Image

        for image_obj in page.images:
            img = Image.open(io.BytesIO(image_obj.data))
            return pytesseract.image_to_string(img, lang="bul+eng")
    except Exception:
        pass
    return ""


def _split_paragraphs(text: str, min_length: int = 20) -> list[str]:
    """Разделя текст на параграфи по двойни нови редове."""
    paragraphs = text.split("\n\n")
    result = []
    for p in paragraphs:
        p = p.replace("\n", " ").strip()
        if len(p) >= min_length:
            result.append(p)
    return result
