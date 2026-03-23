"""
DOCX генератор за експорт на Технически предложения.
Evidence/sources НЕ се включват в exported документа.
"""

from __future__ import annotations

import io
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.models import Project, Generation, TpOutline, ScheduleNormalized


def _set_document_styles(doc) -> None:
    """Задава глобален шрифт, размер и margins на документа."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    import lxml.etree as etree

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # Default font for Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading styles
    heading_colors = {
        1: RGBColor(0x1F, 0x49, 0x7D),  # dark blue
        2: RGBColor(0x2E, 0x74, 0xB5),  # medium blue
        3: RGBColor(0x5B, 0x9B, 0xD5),  # light blue
    }
    for lvl in range(1, 7):
        style_name = f"Heading {lvl}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = "Calibri"
            h.font.size = Pt(max(11, 18 - (lvl - 1) * 2))
            h.font.bold = True
            color = heading_colors.get(lvl, RGBColor(0x40, 0x40, 0x40))
            h.font.color.rgb = color


def _add_page_numbers(doc) -> None:
    """Добавя номерация на страниците в footer-а."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = 2  # RIGHT
        run = p.add_run()
        # Page number field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


async def generate_docx(project_id: str, db: AsyncSession) -> bytes:
    """
    Генерира .docx от заключената структура и генерираните текстове.
    Evidence не се включва в документа.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    project = await db.get(Project, project_id)
    if not project:
        raise ValueError(f"Project {project_id} not found")

    # Взима последния заключен outline
    result = await db.execute(
        select(TpOutline)
        .where(TpOutline.project_id == project_id, TpOutline.status_locked == True)
        .order_by(TpOutline.version.desc())
        .limit(1)
    )
    outline = result.scalar_one_or_none()

    doc = Document()
    doc.core_properties.author = "TP AI"
    doc.core_properties.title = f"Техническо предложение — {project.name}"

    _set_document_styles(doc)
    _add_page_numbers(doc)

    # Заглавие
    title = doc.add_heading(level=1)
    title.clear()
    run = title.add_run("ТЕХНИЧЕСКО ПРЕДЛОЖЕНИЕ")
    run.font.size = Pt(20)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(level=2)
    subtitle.clear()
    run2 = subtitle.add_run(project.name)
    run2.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project.location or project.contracting_authority:
        doc.add_paragraph()
    if project.contracting_authority:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Възложител: ").bold = True
        p.add_run(project.contracting_authority)
    if project.location:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Местоположение: ").bold = True
        p.add_run(project.location)
    if project.tender_date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Дата: ").bold = True
        p.add_run(str(project.tender_date))

    doc.add_page_break()

    if outline:
        sections = outline.outline_json.get(
            "sections", outline.outline_json.get("outline", [])
        )
        await _write_sections(doc, sections, project_id, db, level=2)
    else:
        doc.add_paragraph("⚠ Структурата на ТП не е заключена.")

    # Раздел "ЛИНЕЕН ГРАФИК" от нормализирания график
    sched_result = await db.execute(
        select(ScheduleNormalized)
        .where(ScheduleNormalized.project_id == project_id)
        .order_by(ScheduleNormalized.version.desc())
        .limit(1)
    )
    schedule_norm = sched_result.scalar_one_or_none()
    if schedule_norm:
        doc.add_page_break()
        sched_heading = doc.add_heading(level=2)
        sched_heading.clear()
        sched_heading.add_run("ЛИНЕЕН ГРАФИК").bold = True
        _write_schedule_section(doc, schedule_norm.schedule_json)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _write_schedule_section(doc, schedule_json: dict) -> None:
    """Форматира задачите от нормализирания график в .docx таблица."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn

    tasks = schedule_json.get("tasks", schedule_json.get("normalized", {}).get("tasks", []))
    resources = schedule_json.get("resources", schedule_json.get("normalized", {}).get("resources", []))

    if not tasks:
        doc.add_paragraph("[Графикът не съдържа задачи.]")
        return

    doc.add_paragraph(f"Общо задачи: {len(tasks)}   |   Ресурси: {len(resources)}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Column widths
    widths = [Cm(9), Cm(3.5), Cm(3), Cm(3)]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]

    # Bold header row with shading
    headers = ["Задача", "Продължителност (дни)", "Начало", "Край"]
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        # Gray background for header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tc.get_or_add_tcPr().find(qn("w:shd"))
        from docx.oxml import OxmlElement
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")

    for task in tasks:
        row = table.add_row().cells
        row[0].text = str(task.get("name", "") or "")
        row[1].text = str(task.get("duration_days", task.get("duration", "")) or "")
        row[2].text = str(task.get("start", "") or "")
        row[3].text = str(task.get("finish", task.get("end", "")) or "")
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)


async def _write_sections(
    doc, sections: list, project_id: str, db: AsyncSession, level: int
):
    from sqlalchemy import select
    from docx.shared import Pt
    from app.core.models import Generation

    for section in sections:
        # Support both key names: "section_uid" (old) and "uid" (tender_struct output)
        section_uid = section.get("section_uid") or section.get("uid", "")
        title = section.get("title", "")
        numbering = section.get("display_numbering", "")

        heading = doc.add_heading(level=min(level, 9))
        heading.clear()
        heading.add_run(f"{numbering} {title}".strip()).bold = True

        # Текст от генерацията (ако има)
        result = await db.execute(
            select(Generation)
            .where(
                Generation.project_id == project_id,
                Generation.section_uid == section_uid,
            )
            .order_by(
                Generation.selected.desc(),   # закрепен от потребителя
                Generation.variant.asc(),     # вариант 1 преди 2
                Generation.created_at.desc(),
            )
            .limit(1)
        )
        generation = result.scalar_one_or_none()

        if generation:
            # Evidence НЕ се включва в .docx
            p = doc.add_paragraph(generation.text)
            for run in p.runs:
                run.font.size = Pt(11)
        else:
            p = doc.add_paragraph("[Текстът за тази точка не е генериран.]")
            p.runs[0].font.color.rgb = None  # default gray via style
            p.runs[0].font.italic = True

        # Support both key names: "children" (old) and "subsections" (tender_struct output)
        children = section.get("children") or section.get("subsections", [])
        if children:
            await _write_sections(doc, children, project_id, db, level=level + 1)
